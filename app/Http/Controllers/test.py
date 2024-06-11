import sys

# Получите переданный параметр
file_to_process = sys.argv[1]




try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    # Load the document
    doc = Document(file_to_process)

    # # Get the first section (if exists)
    # if doc.sections:
    #     section = doc.sections[0]

    #     # Get the header for the first section
    #     header = section.header

    #     # Clear the existing content in the header
    #     for paragraph in header.paragraphs:
    #         paragraph.clear()

    #     # Add new text to the header
    #     header.add_paragraph("Новый текст в верхнем колонтитуле")

    for i, section in enumerate(doc.sections[1:]):
        footer = section.footer

        # Clear the existing content in the footer
        for paragraph in footer.paragraphs:
            paragraph.clear()

        # Add a new paragraph with the page number centered
        # paragraph = footer.add_paragraph()
        # paragraph.add_run().add_page_number()
        # paragraph.runs[0].font.size = Pt(12)
        # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save the changes to a new document

        print("Changes saved successfully.")
    else:
        print("The document has no sections.")
    doc.save("test2.docx")
    doc.save(file_to_process)
except Exception as e:
    print(f"An error occurred: {e}")

